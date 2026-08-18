from __future__ import annotations

from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from src.api.core.app import create_app

MT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def docx_bytes(title: str, parties: list[tuple[str, dict[str, str]]], articles: list[tuple[str, str, list[str]]]) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for side, rows in parties:
        doc.add_paragraph(side)
        for key, value in rows.items():
            doc.add_paragraph(f"{key}: {value}")
    for number, heading, paragraphs in articles:
        doc.add_heading(f"Điều {number}. {heading}", level=2)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


CASES = [
    (
        "chia_tai_san.docx",
        "Tôi là Bên A, hãy kiểm tra hợp đồng chia tài sản: tài sản, thanh toán, đăng ký sang tên, rủi ro thiếu căn cứ.",
        docx_bytes(
            "HỢP ĐỒNG THỎA THUẬN CHIA TÀI SẢN",
            [
                ("Bên A", {"Tên": "Nguyễn Văn Minh", "Vai trò": "Người nhận quyền sở hữu căn hộ", "Địa chỉ": "Đà Nẵng"}),
                ("Bên B", {"Tên": "Trần Thị Lan", "Vai trò": "Người nhận khoản thanh toán bù trừ", "Địa chỉ": "Đà Nẵng"}),
            ],
            [
                ("1", "Tài sản phân chia", ["Hai bên thống nhất căn hộ tại phường Hải Châu được giao cho Bên A quản lý, sử dụng và thực hiện thủ tục sang tên.", "Giấy chứng nhận quyền sở hữu ghi số chưa rõ và chưa đính kèm phụ lục hồ sơ pháp lý."]),
                ("2", "Thanh toán bù trừ", ["Bên A thanh toán cho Bên B 800.000.000 đồng trong vòng 30 ngày kể từ ngày ký.", "Nếu chậm thanh toán thì chịu phạt 15% giá trị chậm trả và bồi thường toàn bộ thiệt hại."]),
                ("3", "Thuế, phí và đăng ký", ["Bên A chịu toàn bộ thuế, phí, lệ phí trước bạ và chi phí công chứng."]),
                ("4", "Cam kết", ["Bên B cam kết tài sản không tranh chấp, không bị kê biên, không thế chấp.", "Nếu Điều 8.1 không được thực hiện thì Bên vi phạm chịu trách nhiệm theo Điều 2."]),
            ],
        ),
    ),
    (
        "lao_dong_rui_ro.docx",
        "Tôi là người lao động, hãy kiểm tra hợp đồng lao động này có bất lợi, trái luật, thiếu điều gì không.",
        docx_bytes(
            "HỢP ĐỒNG LAO ĐỘNG RỦI RO",
            [
                ("Người sử dụng lao động", {"Tên": "Công ty Cổ phần TechTop", "Đại diện": "Lê Hoàng", "Chức vụ": "Giám đốc"}),
                ("Người lao động", {"Tên": "Phạm An", "Vai trò": "Kỹ sư AI Agent"}),
            ],
            [
                ("1", "Công việc và địa điểm", ["Người lao động phát triển AI agent quản lý hợp đồng, vận hành hệ thống và hỗ trợ khách hàng theo yêu cầu của công ty."]),
                ("2", "Thời giờ làm việc", ["Người lao động làm 12 giờ/ngày, tăng ca theo yêu cầu dự án và không tính lương làm thêm.", "Người lao động không có ngày nghỉ trong tháng cao điểm."]),
                ("3", "Tiền lương", ["Lương cơ bản 12.000.000 đồng/tháng, thưởng dự án do công ty tự quyết định.", "Nếu nghỉ trước 24 tháng thì người lao động hoàn trả 100% chi phí đào tạo và 03 tháng lương."]),
                ("4", "Giữ hồ sơ", ["Người sử dụng lao động giữ bản gốc bằng cấp và căn cước công dân đến khi kết thúc hợp đồng."]),
                ("5", "Chấm dứt", ["Công ty có thể chấm dứt ngay nếu kết quả công việc không đạt kỳ vọng nội bộ.", "Nội dung khác thực hiện theo Điều 9.9 của hợp đồng này."]),
            ],
        ),
    ),
    (
        "ba_ben_bao_lanh_ai.docx",
        "Tôi là Bên A, hãy kiểm tra quan hệ nhiều bên, nghĩa vụ bảo lãnh, nghiệm thu và rủi ro pháp lý.",
        docx_bytes(
            "HỢP ĐỒNG BA BÊN TRIỂN KHAI AI AGENT",
            [
                ("Bên A", {"Tên": "GDGoC-DUT", "Vai trò": "Khách hàng sử dụng hệ thống"}),
                ("Bên B", {"Tên": "Công ty Cổ phần TechTop", "Vai trò": "Đơn vị phát triển AI agent"}),
                ("Bên C", {"Tên": "Quỹ Hỗ trợ Sáng tạo Miền Trung", "Vai trò": "Bên bảo lãnh thanh toán một phần"}),
            ],
            [
                ("1", "Phạm vi", ["Bên B cung cấp cho Bên A AI agent quản lý hợp đồng, dashboard rủi ro và báo cáo định kỳ.", "Bên C bảo lãnh cho Bên A tối đa 200.000.000 đồng nếu Bên A chậm thanh toán vì thủ tục giải ngân."]),
                ("2", "Thanh toán", ["Bên A thanh toán cho Bên B 500.000.000 đồng theo 3 đợt; Bên C chuyển phần bảo lãnh trực tiếp cho Bên B sau khi nhận biên bản nghiệm thu."]),
                ("3", "Nghiệm thu", ["Bên B bàn giao cho Bên A mã nguồn, tài liệu API và tài khoản quản trị trong 45 ngày.", "Nếu Bên A không phản hồi trong 03 ngày làm việc thì sản phẩm được xem là đã nghiệm thu."]),
                ("4", "Sở hữu trí tuệ và dữ liệu", ["Bên B giữ quyền sở hữu nền tảng lõi; Bên A sở hữu dữ liệu nội bộ và báo cáo sinh ra từ dữ liệu của Bên A.", "Bên B được dùng dữ liệu đã ẩn danh để cải tiến mô hình nếu không có phản đối bằng văn bản."]),
                ("5", "Vi phạm", ["Mức phạt vi phạm là 20% tổng giá trị hợp đồng. Các trách nhiệm tại Điều 7.4 vẫn áp dụng dù hợp đồng chấm dứt."]),
            ],
        ),
    ),
]


def main() -> None:
    client = TestClient(create_app())
    for filename, question, data in CASES:
        response = client.post("/api/v1/contracts/analyze-docx", files={"file": (filename, data, MT)}, data={"question": question})
        body = response.json()
        a = body["module_a"]
        b = body["module_b"]
        print("\n--- INPUT ---")
        print(filename)
        print(question)
        print("--- OUTPUT A ---")
        print("parties=", [(p["side"], p["name"], p["role"]) for p in a["parties"]])
        print("clauses=", len(a["clauses"]), "obligations=", len(a["obligations"]))
        print("relationships=", [(r["from_party"], r["relation"], r["to_party"], r["source_clause_id"]) for r in a["relationships"]])
        print("bad_refs=", [(r["target_label"], r["source_clause_id"]) for r in a["internal_references"] if not r["target_exists"]])
        print("risks=", [(r["kind"], r["severity"], r["source_clause_id"]) for r in a["risk_candidates"]])
        print("--- OUTPUT B ---")
        print("profile=", b["question_profile"])
        print("coverage=", b["coverage"])
        print("legal_plan=", [(item.get("topic"), item.get("priority") or item.get("reason")) for item in b["legal_search_plan"]])

    filename, question, data = CASES[1]
    print("\n=== LLM/RAG CHECK on lao_dong_rui_ro.docx ===")
    response = client.post("/api/v1/contracts/analyze-docx-llm", files={"file": (filename, data, MT)}, data={"question": question, "enable_rag": "true"})
    body = response.json()
    rag = body.get("rag_evidence") or {}
    review = body.get("llm_review") or {}
    print("http=", response.status_code)
    print("rag_queries=", rag.get("queries"))
    print("rag_n=", len(rag.get("evidence") or []), "rag_warnings=", rag.get("warnings"))
    print("llm_status=", review.get("status"))
    print("llm_warnings=", review.get("warnings"))
    print("prompt_preview=", (review.get("prompt_preview") or "")[:1200].replace("\n", " | "))
    print("answer=", (review.get("answer_markdown") or "")[:2400].replace("\n", " | "))
    print("usage=", review.get("usage"))


if __name__ == "__main__":
    main()
