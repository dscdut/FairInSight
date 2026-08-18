from __future__ import annotations

from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from src.api.core.app import create_app
from src.services.contracts import analyze_contract_docx_bytes, extract_contract_docx_bytes


ROOT = Path(__file__).resolve().parents[3]
SAMPLE_DOCX = ROOT / ".ask" / "generated_contracts" / "hop_dong_gdgoc_dut_techtop_ai_agent.docx"


def _sample_bytes() -> bytes:
    assert SAMPLE_DOCX.exists(), f"Missing sample DOCX: {SAMPLE_DOCX}"
    return SAMPLE_DOCX.read_bytes()


def _docx_bytes_from_document(document: Document) -> bytes:
    from io import BytesIO

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_contract_module_a_extracts_clean_contract_data() -> None:
    result = extract_contract_docx_bytes(_sample_bytes(), filename=SAMPLE_DOCX.name)

    assert result.document_info["filename"] == SAMPLE_DOCX.name
    assert result.document_info["clause_count"] >= 13
    assert len(result.parties) == 2
    assert {party.name for party in result.parties} >= {"GDGoC-DUT", "Công ty Cổ phần TechTop"}
    assert len(result.tables) >= 5
    assert any(item.actor == "Bên B" for item in result.obligations)
    assert any(item["value"] == "120.000.000 VNĐ" for item in result.money_terms)
    assert any(ref.target_label == "Điều 1.2" and ref.target_exists for ref in result.internal_references)
    assert result.clean_context["summary"]["risk_count"] >= 0


def test_contract_module_a_extracts_three_parties_and_relationships() -> None:
    doc = Document()
    doc.add_paragraph("HỢP ĐỒNG TRIỂN KHAI NỀN TẢNG AI BA BÊN")
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Nội dung"
    table.rows[0].cells[1].text = "Bên A"
    table.rows[0].cells[2].text = "Bên B"
    table.rows[0].cells[3].text = "Bên C"
    for label, a, b, c in [
        ("Tên đơn vị", "Công ty Alpha", "Công ty Beta", "Ngân hàng Gamma"),
        ("Vai trò", "Bên đặt hàng", "Bên triển khai", "Bên bảo lãnh thanh toán"),
        ("Đại diện", "Ông A", "Bà B", "Ông C"),
        ("Email", "a@example.test", "b@example.test", "c@example.test"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = label, a, b, c
    doc.add_paragraph("Điều 1. Phạm vi")
    doc.add_paragraph("1.1. Bên B cung cấp hệ thống AI quản lý cho Bên A theo yêu cầu tại Điều 2.1.")
    doc.add_paragraph("1.2. Bên C bảo lãnh cho Bên A nghĩa vụ thanh toán tại Điều 3.1.")
    doc.add_paragraph("Điều 2. Dữ liệu và phối hợp")
    doc.add_paragraph("2.1. Bên A cung cấp dữ liệu cho Bên B trong vòng 05 ngày làm việc.")
    doc.add_paragraph("2.2. Bên C phối hợp với Bên B để xác nhận hạn mức bảo lãnh.")
    doc.add_paragraph("Điều 3. Thanh toán")
    doc.add_paragraph("3.1. Bên A thanh toán cho Bên B 200.000.000 VNĐ sau nghiệm thu.")

    result = extract_contract_docx_bytes(_docx_bytes_from_document(doc), filename="three-party.docx")

    assert [party.side for party in result.parties] == ["Bên A", "Bên B", "Bên C"]
    assert {party.name for party in result.parties} == {"Công ty Alpha", "Công ty Beta", "Ngân hàng Gamma"}
    assert {rel.relation for rel in result.relationships} >= {"cung cấp cho", "bảo lãnh cho", "thanh toán cho", "phối hợp với"}
    assert any(ref.target_label == "Điều 2.1" and ref.target_id == "D2.1" for ref in result.internal_references)
    assert result.clean_context["summary"]["relationship_count"] >= 4
    assert not result.warnings


def test_contract_module_a_detects_multiple_a_sides_and_broken_reference() -> None:
    doc = Document()
    doc.add_paragraph("HỢP ĐỒNG HỢP TÁC NHIỀU BÊN A")
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Nội dung"
    table.rows[0].cells[1].text = "Bên A1"
    table.rows[0].cells[2].text = "Bên A2"
    table.rows[0].cells[3].text = "Bên B"
    for label, a1, a2, b in [
        ("Tên đơn vị", "Đơn vị sở hữu dữ liệu", "Đơn vị vận hành", "Nhà cung cấp AI"),
        ("Vai trò", "Cung cấp dữ liệu", "Nghiệm thu vận hành", "Phát triển hệ thống"),
        ("Đại diện", "Ông A1", "Bà A2", "Ông B"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = label, a1, a2, b
    doc.add_paragraph("Điều 1. Phân công")
    doc.add_paragraph("1.1. Bên A1 cung cấp dữ liệu cho Bên B trong vòng 03 ngày làm việc.")
    doc.add_paragraph("1.2. Bên B bàn giao cho Bên A2 bản demo theo Điều 9.9.")
    doc.add_paragraph("Điều 2. Nghiệm thu")
    doc.add_paragraph("2.1. Bên A2 xác nhận nghiệm thu trong vòng 07 ngày làm việc.")

    result = extract_contract_docx_bytes(_docx_bytes_from_document(doc), filename="multi-a.docx")

    assert [party.side for party in result.parties] == ["Bên A1", "Bên A2", "Bên B"]
    assert any(rel.from_party == "Bên A1" and rel.to_party == "Bên B" for rel in result.relationships)
    assert any(rel.from_party == "Bên B" and rel.to_party == "Bên A2" for rel in result.relationships)
    broken = [ref for ref in result.internal_references if not ref.target_exists]
    assert any(ref.target_label == "Điều 9.9" for ref in broken)
    assert any(risk.kind == "broken_reference" for risk in result.risk_candidates)


def test_contract_module_a_handles_labor_contract_risk_terms() -> None:
    doc = Document()
    doc.add_paragraph("HỢP ĐỒNG LAO ĐỘNG")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Nội dung"
    table.rows[0].cells[1].text = "Người sử dụng lao động"
    table.rows[0].cells[2].text = "Người lao động"
    for label, employer, worker in [
        ("Tên đơn vị", "Công ty May Sao Đêm", "Nguyễn Văn Nam"),
        ("Vai trò", "Bên thuê lao động", "Nhân viên vận hành"),
        ("Đại diện", "Bà Trần Quản Lý", "Ông Nguyễn Văn Nam"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = label, employer, worker
    doc.add_paragraph("Điều 1. Công việc")
    doc.add_paragraph("1.1. Người lao động vận hành chuyền may theo phân công của Người sử dụng lao động.")
    doc.add_paragraph("Điều 2. Thời giờ làm việc")
    doc.add_paragraph("2.1. Người lao động làm việc 12 giờ/ngày, tăng ca theo yêu cầu sản xuất và không tính lương làm thêm.")
    doc.add_paragraph("2.2. Người lao động không có ngày nghỉ trong tháng cao điểm.")
    doc.add_paragraph("Điều 3. Giấy tờ")
    doc.add_paragraph("3.1. Người sử dụng lao động giữ bản gốc bằng cấp của Người lao động đến khi kết thúc hợp đồng.")

    result = extract_contract_docx_bytes(_docx_bytes_from_document(doc), filename="labor-risk.docx")

    assert [party.side for party in result.parties] == ["Người sử dụng lao động", "Người lao động"]
    assert any(item.actor == "Người lao động" for item in result.obligations)
    assert {risk.kind for risk in result.risk_candidates} >= {
        "labor_overtime",
        "labor_rest",
        "labor_document_retention",
    }
    assert result.clean_context["summary"]["party_count"] == 2


def test_contract_module_b_loops_until_coverage_and_builds_legal_plan() -> None:
    result = analyze_contract_docx_bytes(
        _sample_bytes(),
        filename=SAMPLE_DOCX.name,
        question="Kiểm tra rủi ro và căn cứ pháp luật Việt Nam về thanh toán, nghiệm thu, bảo mật.",
    )

    assert result.module_b is not None
    assert result.module_b.reasoning_loops
    assert result.module_b.selected_clause_ids
    assert result.module_b.coverage["selected_count"] > 0
    assert result.module_b.legal_search_plan
    assert any(task["topic"] == "payment" for task in result.module_b.legal_search_plan)
    assert result.module_c is not None
    assert "Kết quả kiểm tra hợp đồng" in result.module_c.report_markdown
    assert "`D4.2`" in result.module_c.report_markdown
    assert "`D5.4`" in result.module_c.report_markdown


def test_contract_module_b_handles_vietnamese_case_and_ascii_question() -> None:
    questions = [
        "Kiểm tra rủi ro và căn cứ pháp luật Việt Nam về thanh toán, nghiệm thu, bảo mật.",
        "kiem tra rui ro va can cu phap luat ve thanh toan nghiem thu bao mat",
        "KIỂM TRA ĐIỀU 4.1 VÀ ĐIỀU 9 VỀ THANH TOÁN, BẢO MẬT",
    ]

    for question in questions:
        result = analyze_contract_docx_bytes(_sample_bytes(), filename=SAMPLE_DOCX.name, question=question)

        assert result.module_b is not None
        assert result.module_b.selected_clause_ids
        assert "payment" in {
            topic
            for step in result.module_b.reasoning_loops
            for topic in step["coverage"]["topics"]
        }


def test_contract_module_b_keeps_utf8_intents_and_legal_topics() -> None:
    result = analyze_contract_docx_bytes(
        _sample_bytes(),
        filename=SAMPLE_DOCX.name,
        question="Kiểm tra rủi ro và căn cứ pháp luật Việt Nam về thanh toán, nghiệm thu, bảo mật, sở hữu trí tuệ.",
    )

    assert result.module_b is not None
    assert set(result.module_b.question_profile["intents"]) >= {"legal_check", "risk_review"}
    assert set(result.module_b.question_profile["requested_topics"]) >= {
        "payment",
        "acceptance_schedule",
        "confidentiality_data",
        "intellectual_property",
    }
    assert result.module_b.coverage["ready"] is True
    assert {
        task["topic"] for task in result.module_b.legal_search_plan
    } >= {
        "payment",
        "acceptance_schedule",
        "confidentiality_data",
        "intellectual_property",
    }


def test_contract_docx_api_is_separate_from_chat_route() -> None:
    client = TestClient(create_app())
    with SAMPLE_DOCX.open("rb") as handle:
        response = client.post(
            "/api/v1/contracts/extract-docx",
            files={
                "file": (
                    SAMPLE_DOCX.name,
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    body = response.json()
    assert body["module"] == "A"
    assert len(body["clauses"]) >= 13
    assert body["parties"][0]["name"] == "GDGoC-DUT"
